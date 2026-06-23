"""10-K Word (.docx) 上传与正文抽取。"""
import io
from datetime import datetime
from typing import Any, Dict

from flask import current_app

from app.services.financial_reports import save_report_supplement


def extract_text_from_docx(data: bytes) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，请执行 pip install python-docx") from exc

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return {"text": text, "char_count": len(text)}


def save_uploaded_docx_supplement(report_id: int, file_storage) -> Dict[str, Any]:
    max_bytes = int(current_app.config.get("FINANCIAL_PDF_MAX_BYTES", 50 * 1024 * 1024))
    filename = file_storage.filename or "supplement.docx"
    if not filename.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 文件")
    data = file_storage.read()
    if len(data) > max_bytes:
        raise ValueError(f"文件超过 {max_bytes // (1024 * 1024)}MB 限制")
    if len(data) < 100:
        raise ValueError("文件过小，不是有效的 Word 文档")

    extracted = extract_text_from_docx(data)
    if extracted["char_count"] < 200:
        raise ValueError("Word 正文过少（可能为扫描版），无法提取有效文字")

    meta = {
        "uploaded_at": datetime.now().isoformat(timespec="seconds"),
        "char_count": extracted["char_count"],
        "form_type": "10-K",
        "filename": filename,
    }
    save_report_supplement(report_id, data, filename, extracted["text"], meta)
    return extracted
